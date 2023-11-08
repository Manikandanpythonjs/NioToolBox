from tkinter import PhotoImage
import customtkinter
from PIL import Image
from CTkMessagebox import CTkMessagebox
from tkinter import filedialog, StringVar
from docx import Document

# doc = aw.Document("Input.pdf")
# doc.save("Output.txt")


class WordtoText(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        self.selectedpath = StringVar()

        def center_window(self, width, height):
            screen_width = self.winfo_screenwidth()
            screen_height = self.winfo_screenheight()
            x_coordinate = (screen_width - width) // 2
            y_coordinate = (screen_height - height) // 2
            self.geometry(f"{width}x{height}+{x_coordinate}+{y_coordinate}")

        self.title("wordtotext")
        self._set_appearance_mode("light")
        center_window(self, 300, 200)
        self.resizable(False, False)
        self.configure(fg_color="white")
        self.attributes('-topmost', True)

        self.logotitle = customtkinter.CTkLabel(
            self, text="WORD TO TEXT", text_color="grey")
        self.logotitle.pack(padx=10, pady=10, anchor="n")

        self.openfile = customtkinter.CTkButton(
            self, text="Change WORD to TEXT ", command=self.open_file)
        self.openfile.pack(padx=20, pady=20)

        self.selectedpathLabel = customtkinter.CTkLabel(
            self,  text=" ", text_color="grey")
        self.selectedpathLabel.pack(padx=10, pady=10, anchor="n")

    def open_file(self):
        self.attributes('-topmost', False)
        self.file_path = filedialog.askopenfilename(
            filetypes=[("Word Files", "*.docx")])
        if self.file_path:

            self.selectedpathLabel.configure(text=self.file_path)
            self.openfile.configure(
                text="Convert to Text", command=self.converttexttoword)
            self.selectedpath.set(self.file_path)
            print("Selected file:", self.file_path)
        self.attributes('-topmost', True)

    def converttexttoword(self):
        self.wm_attributes("-topmost", False)
        if (self.file_path):

            self.selectedpathLabel.configure(text=self.file_path)

            self.openfile.configure(
                text="Please Wait ....")


            self.output_word_file = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Text Files", "*.txt")])
            
            try:
            
                doc = Document(self.file_path)
                text = ""

                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

                # Save the text to the output text file
                with open(self.output_word_file, "w", encoding="utf-8") as file:
                    file.write(text)
            except Exception as e:
                            CTkMessagebox(
                            title="Error", message=str(e), icon="cancel")
          
            self.wm_attributes("-topmost", True)
            self.openfile.configure(
                text="Change Word to Text", command=self.open_file)

            self.selectedpath.set("")

        else:

            pass
